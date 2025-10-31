# election_project/election_app/views.py
from django.shortcuts import render, redirect, get_object_or_404
from django.views import View
from django.contrib import messages
from django.contrib.auth import authenticate, login as django_login, logout as django_logout
from django.contrib.auth.models import User
from django.utils import timezone
from django.db.models import Sum, Count
from django.http import HttpResponse, JsonResponse
from .models import ElectionSetting, Candidate, Voter, ActivityLog
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Paragraph, Table, TableStyle
from reportlab.lib import colors
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
import datetime

class IndexView(View):
    def get(self, request):
        candidates = Candidate.objects.all()
        total_votes = candidates.aggregate(total=Sum('votes'))['total'] or 0
        for candidate in candidates:
            candidate.percentage = round((candidate.votes / total_votes * 100) if total_votes > 0 else 0, 2)
        return render(request, 'index.html', {'candidates': candidates})

    def post(self, request):
        role = request.POST.get('role')
        password = request.POST.get('password')
        if role == 'voter':
            dept_id = request.POST.get('dept_id')
            try:
                voter = Voter.objects.get(dept_id=dept_id)
                if voter.check_password(password):
                    if voter.has_voted:
                        messages.error(request, 'You have already voted')
                    else:
                        request.session['voter_id'] = voter.id
                        return redirect('voter_dashboard')
                else:
                    messages.error(request, 'Invalid credentials')
            except Voter.DoesNotExist:
                messages.error(request, 'Invalid credentials')
        elif role == 'admin':
            username = request.POST.get('username')
            user = authenticate(request, username=username, password=password)
            if user is not None:
                django_login(request, user)
                return redirect('dashboard')
            else:
                messages.error(request, 'Invalid credentials')
        # Reload index with messages
        return self.get(request)

class SetupAdminView(View):
    def get(self, request):
        if User.objects.exists():
            return redirect('login')
        return render(request, 'setup_admin.html')

    def post(self, request):
        if User.objects.exists():
            return redirect('login')
        username = request.POST.get('username')
        password = request.POST.get('password')
        if username and password:
            user = User.objects.create_user(username=username, password=password, is_staff=True, is_superuser=True)
            messages.success(request, 'Admin setup complete. Please login.')
            return redirect('login')
        messages.error(request, 'Please fill in both fields')
        return render(request, 'setup_admin.html')

class LoginView(View):
    def get(self, request):
        if not User.objects.exists():
            return redirect('setup_admin')
        return render(request, 'login.html')

    def post(self, request):
        role = request.POST.get('role')
        password = request.POST.get('password')
        if role == 'voter':
            dept_id = request.POST.get('dept_id')
            try:
                voter = Voter.objects.get(dept_id=dept_id)
                if voter.check_password(password):
                    if voter.has_voted:
                        messages.error(request, 'You have already voted')
                    else:
                        request.session['voter_id'] = voter.id
                        return redirect('voter_dashboard')
                else:
                    messages.error(request, 'Invalid credentials')
            except Voter.DoesNotExist:
                messages.error(request, 'Invalid credentials')
        elif role == 'admin':
            username = request.POST.get('username')
            user = authenticate(request, username=username, password=password)
            if user is not None:
                django_login(request, user)
                return redirect('dashboard')
            else:
                messages.error(request, 'Invalid credentials')
        return render(request, 'login.html')

class LogoutView(View):
    def get(self, request):
        django_logout(request)
        request.session.flush()
        return redirect('login')

class DashboardView(View):
    def get(self, request):
        if not request.user.is_authenticated:
            return redirect('login')
        settings = ElectionSetting.objects.first() or ElectionSetting.objects.create()
        candidates = Candidate.objects.all()
        voters = Voter.objects.all()
        activities = ActivityLog.objects.order_by('-time')[:4]
        total_candidates = candidates.count()
        total_voters = voters.count()
        votes_cast = candidates.aggregate(total=Sum('votes'))['total'] or 0
        participation = round((votes_cast / total_voters * 100) if total_voters > 0 else 0)
        election_status = self.get_election_status(settings)
        # Update candidate statuses
        if votes_cast > 0:
            max_votes = candidates.aggregate(max=Sum('votes'))['max']
            for c in candidates:
                c.status = 'Leading' if c.votes == max_votes else 'Running'
                c.status_color = 'green' if c.votes == max_votes else 'blue'
                c.save()
        context = {
            'settings': settings,
            'total_candidates': total_candidates,
            'total_voters': total_voters,
            'votes_cast': votes_cast,
            'participation': participation,
            'activities': activities,
            'candidates': candidates,
            'election_status': election_status,
            'notifications': 3,  # Placeholder
        }
        return render(request, 'dashboard.html', context)

    def get_election_status(self, settings):
        now = timezone.now()
        if now < settings.start_date:
            return {'status': 'Not Started', 'label': 'Time to Start', 'target': settings.start_date.isoformat(), 'started': False, 'ended': False}
        elif now <= settings.end_date:
            return {'status': 'Ongoing', 'label': 'Time Remaining', 'target': settings.end_date.isoformat(), 'started': True, 'ended': False}
        else:
            return {'status': 'Ended', 'label': 'Election Status', 'target': None, 'started': True, 'ended': True}

class CandidatesView(View):
    def get(self, request):
        if not request.user.is_authenticated:
            return redirect('login')
        candidates = Candidate.objects.all()
        total_votes = Candidate.objects.aggregate(total=Sum('votes'))['total'] or 0
        context = {'candidates': candidates, 'total_votes': total_votes}
        return render(request, 'candidates.html', context)

class AddCandidateView(View):
    def post(self, request):
        if not request.user.is_authenticated:
            return redirect('login')
        name = request.POST.get('candidate-name')
        department = request.POST.get('candidate-department')
        position = request.POST.get('candidate-position')
        photo = request.FILES.get('candidate-photo')
        party_photo = request.FILES.get('candidate-party-photo')
        candidate = Candidate(name=name, department=department, position=position, photo=photo, party_photo=party_photo)
        candidate.save()
        ActivityLog.objects.create(type='Candidate added', description=f'{name} added as candidate', icon='fa-user-tie', color='purple')
        messages.success(request, 'Candidate added successfully')
        return redirect('candidates')

class EditCandidateView(View):
    def post(self, request, pk):
        if not request.user.is_authenticated:
            return redirect('login')
        candidate = get_object_or_404(Candidate, pk=pk)
        candidate.name = request.POST.get('candidate-name', candidate.name)
        candidate.department = request.POST.get('candidate-department', candidate.department)
        candidate.position = request.POST.get('candidate-position', candidate.position)
        if 'candidate-photo' in request.FILES:
            candidate.photo = request.FILES['candidate-photo']
        if 'candidate-party-photo' in request.FILES:
            candidate.party_photo = request.FILES['candidate-party-photo']
        candidate.save()
        ActivityLog.objects.create(type='Candidate updated', description=f'{candidate.name} details updated', icon='fa-edit', color='blue')
        messages.success(request, 'Candidate updated successfully')
        return redirect('candidates')

class DeleteCandidateView(View):
    def post(self, request, pk):
        if not request.user.is_authenticated:
            return redirect('login')
        candidate = get_object_or_404(Candidate, pk=pk)
        name = candidate.name
        candidate.delete()
        ActivityLog.objects.create(type='Candidate removed', description=f'{name} removed from candidates', icon='fa-trash', color='red')
        messages.success(request, 'Candidate deleted successfully')
        return redirect('candidates')

class VotersView(View):
    def get(self, request):
        if not request.user.is_authenticated:
            return redirect('login')
        voters = Voter.objects.all()
        return render(request, 'voters.html', {'voters': voters})

class AddVoterView(View):
    def post(self, request):
        if not request.user.is_authenticated:
            return redirect('login')
        name = request.POST.get('voter-name')
        sex = request.POST.get('voter-sex')
        status = request.POST.get('voter-status')
        major_minor = request.POST.get('voter-major-minor')
        department = request.POST.get('voter-department')
        dept_id = request.POST.get('voter-dept-id')
        password = request.POST.get('voter-password')
        photo = request.FILES.get('voter-photo')
        voter = Voter(name=name, sex=sex, status=status, major_minor=major_minor, department=department, dept_id=dept_id, photo=photo)
        voter.set_password(password)
        ActivityLog.objects.create(type='New voter registered', description=f'{name} just registered to vote', icon='fa-user-plus', color='green')
        messages.success(request, 'Voter added successfully')
        return redirect('voters')

class EditVoterView(View):
    def post(self, request, pk):
        if not request.user.is_authenticated:
            return redirect('login')
        voter = get_object_or_404(Voter, pk=pk)
        voter.name = request.POST.get('voter-name', voter.name)
        voter.sex = request.POST.get('voter-sex', voter.sex)
        voter.status = request.POST.get('voter-status', voter.status)
        voter.major_minor = request.POST.get('voter-major-minor', voter.major_minor)
        voter.department = request.POST.get('voter-department', voter.department)
        voter.dept_id = request.POST.get('voter-dept-id', voter.dept_id)
        if 'voter-password' in request.POST and request.POST['voter-password']:
            voter.set_password(request.POST['voter-password'])
        if 'voter-photo' in request.FILES:
            voter.photo = request.FILES['voter-photo']
        voter.save()
        ActivityLog.objects.create(type='Voter updated', description=f'{voter.name} details updated', icon='fa-edit', color='blue')
        messages.success(request, 'Voter updated successfully')
        return redirect('voters')

class DeleteVoterView(View):
    def post(self, request, pk):
        if not request.user.is_authenticated:
            return redirect('login')
        voter = get_object_or_404(Voter, pk=pk)
        name = voter.name
        voter.delete()
        ActivityLog.objects.create(type='Voter removed', description=f'{name} removed from voters', icon='fa-trash', color='red')
        messages.success(request, 'Voter deleted successfully')
        return redirect('voters')

class ResultsView(View):
    def get(self, request):
        if not request.user.is_authenticated:
            return redirect('login')
        candidates = Candidate.objects.all()
        year_counts = Voter.objects.filter(has_voted=True).values('status').annotate(count=Count('status'))
        major_counts = Voter.objects.filter(has_voted=True).values('major_minor').annotate(count=Count('major_minor'))
        total_voted = Voter.objects.filter(has_voted=True).count()
        year_percent = {item['status']: round(item['count'] / total_voted * 100, 1) if total_voted > 0 else 0 for item in year_counts}
        major_percent = {item['major_minor']: round(item['count'] / total_voted * 100, 1) if total_voted > 0 else 0 for item in major_counts}
        context = {
            'candidates': candidates,
            'year_percent': year_percent,
            'major_percent': major_percent,
        }
        return render(request, 'results.html', context)

class ResultsJSONView(View):
    def get(self, request):
        if not request.user.is_authenticated:
            return JsonResponse({'error': 'Unauthorized'}, status=401)
        candidates = list(Candidate.objects.values('name', 'votes'))
        year_counts = list(Voter.objects.filter(has_voted=True).values('status').annotate(count=Count('status')))
        major_counts = list(Voter.objects.filter(has_voted=True).values('major_minor').annotate(count=Count('major_minor')))
        total_voted = Voter.objects.filter(has_voted=True).count()
        year_percent = {item['status']: round(item['count'] / total_voted * 100, 1) if total_voted > 0 else 0 for item in year_counts}
        major_percent = {item['major_minor']: round(item['count'] / total_voted * 100, 1) if total_voted > 0 else 0 for item in major_counts}
        return JsonResponse({
            'candidates': candidates,
            'year_percent': year_percent,
            'major_percent': major_percent,
        })

class DownloadPDFView(View):
    def get(self, request):
        if not request.user.is_authenticated:
            return redirect('login')
        candidates = Candidate.objects.all()
        year_counts = Voter.objects.filter(has_voted=True).values('status').annotate(count=Count('status'))
        major_counts = Voter.objects.filter(has_voted=True).values('major_minor').annotate(count=Count('major_minor'))
        total_voted = Voter.objects.filter(has_voted=True).count()
        year_percent = {item['status']: round(item['count'] / total_voted * 100, 1) if total_voted > 0 else 0 for item in year_counts}
        major_percent = {item['major_minor']: round(item['count'] / total_voted * 100, 1) if total_voted > 0 else 0 for item in major_counts}

        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter)
        styles = getSampleStyleSheet()
        story = []

        story.append(Paragraph('Election Results Report', styles['Title']))
        story.append(Paragraph(f'Generated on: {datetime.date.today()}', styles['Normal']))
        story.append(Paragraph('', styles['Normal']))  # Spacer

        story.append(Paragraph('Candidates Results', styles['Heading2']))
        candidate_data = [['Name', 'Votes']]
        for candidate in candidates:
            candidate_data.append([candidate.name, candidate.votes])
        candidate_table = Table(candidate_data)
        candidate_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        story.append(candidate_table)
        story.append(Paragraph('', styles['Normal']))  # Spacer

        story.append(Paragraph('Votes by Year Level (%)', styles['Heading2']))
        year_data = [['Year Level', 'Percentage']]
        for key, value in year_percent.items():
            year_data.append([key, f'{value}%'])
        year_table = Table(year_data)
        year_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        story.append(year_table)
        story.append(Paragraph('', styles['Normal']))  # Spacer

        story.append(Paragraph('Votes by Major/Minor (%)', styles['Heading2']))
        major_data = [['Type', 'Percentage']]
        for key, value in major_percent.items():
            major_data.append([key, f'{value}%'])
        major_table = Table(major_data)
        major_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        story.append(major_table)

        doc.build(story)
        buffer.seek(0)
        response = HttpResponse(buffer, content_type='application/pdf')
        response['Content-Disposition'] = 'attachment; filename="election_results.pdf"'
        return response

class DownloadWordView(View):
    def get(self, request):
        if not request.user.is_authenticated:
            return redirect('login')
        candidates = Candidate.objects.all()
        year_counts = Voter.objects.filter(has_voted=True).values('status').annotate(count=Count('status'))
        major_counts = Voter.objects.filter(has_voted=True).values('major_minor').annotate(count=Count('major_minor'))
        total_voted = Voter.objects.filter(has_voted=True).count()
        year_percent = {item['status']: round(item['count'] / total_voted * 100, 1) if total_voted > 0 else 0 for item in year_counts}
        major_percent = {item['major_minor']: round(item['count'] / total_voted * 100, 1) if total_voted > 0 else 0 for item in major_counts}

        doc = Document()
        doc.add_heading('Election Results Report', level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(f'Generated on: {datetime.date.today()}')

        doc.add_heading('Candidates Results', level=2)
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Name'
        hdr_cells[1].text = 'Votes'
        for candidate in candidates:
            row_cells = table.add_row().cells
            row_cells[0].text = candidate.name
            row_cells[1].text = str(candidate.votes)

        doc.add_heading('Votes by Year Level (%)', level=2)
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Year Level'
        hdr_cells[1].text = 'Percentage'
        for key, value in year_percent.items():
            row_cells = table.add_row().cells
            row_cells[0].text = key
            row_cells[1].text = f'{value}%'

        doc.add_heading('Votes by Major/Minor (%)', level=2)
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Type'
        hdr_cells[1].text = 'Percentage'
        for key, value in major_percent.items():
            row_cells = table.add_row().cells
            row_cells[0].text = key
            row_cells[1].text = f'{value}%'

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        response = HttpResponse(buffer, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename="election_results.docx"'
        return response

class SettingsView(View):
    def get(self, request):
        if not request.user.is_authenticated:
            return redirect('login')
        settings = ElectionSetting.objects.first() or ElectionSetting.objects.create()
        return render(request, 'settings.html', {'settings': settings})

    def post(self, request):
        settings = ElectionSetting.objects.first() or ElectionSetting.objects.create()
        settings.admin_name = request.POST.get('admin-name-input', settings.admin_name)
        settings.admin_role = request.POST.get('admin-role-input', settings.admin_role)
        if 'admin-photo' in request.FILES:
            settings.admin_avatar = request.FILES['admin-photo']
        settings.start_date = request.POST.get('election-start-date')
        settings.end_date = request.POST.get('election-end-date')
        settings.save()
        messages.success(request, 'Settings saved successfully')
        return redirect('settings')

class VoterDashboardView(View):
    def get(self, request):
        voter_id = request.session.get('voter_id')
        if not voter_id:
            return redirect('login')
        voter = get_object_or_404(Voter, id=voter_id)
        settings = ElectionSetting.objects.first()
        now = timezone.now()
        if now < settings.start_date or now > settings.end_date:
            messages.error(request, 'Election is not ongoing')
            del request.session['voter_id']
            return redirect('login')
        if voter.has_voted:
            messages.error(request, 'You have already voted')
            del request.session['voter_id']
            return redirect('login')
        candidates = Candidate.objects.all()
        return render(request, 'voter_dashboard.html', {'candidates': candidates})

    def post(self, request):
        voter_id = request.session.get('voter_id')
        if not voter_id:
            return redirect('login')
        voter = get_object_or_404(Voter, id=voter_id)
        candidate_id = request.POST.get('candidate')
        if not candidate_id:
            messages.error(request, 'Please select a candidate')
            return self.get(request)
        candidate = get_object_or_404(Candidate, id=candidate_id)
        candidate.votes += 1
        candidate.save()
        voter.has_voted = True
        voter.save()
        ActivityLog.objects.create(type='Vote recorded', description=f'Vote cast for {candidate.name} by {voter.name}', icon='fa-vote-yea', color='blue')
        del request.session['voter_id']
        return render(request, 'voter_dashboard.html', {'success': True})